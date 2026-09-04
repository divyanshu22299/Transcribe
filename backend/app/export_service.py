import json
import csv
import io
from pathlib import Path
from typing import List, Dict, Any, Optional
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

from app.models import TranscriptionResult, Segment, AudioAnalysis
from app.config import EXPORTS_DIR


def export_to_csv(result: TranscriptionResult, delimiter: str = ",") -> str:
    """Generate CSV or TSV string representation."""
    output = io.StringIO()
    writer = csv.writer(output, delimiter=delimiter)
    
    # Headers
    headers = [
        "Audio Filename",
        "Segment ID",
        "Speaker",
        "Gender",
        "Start Time (s)",
        "End Time (s)",
        "Start Time (Formatted)",
        "End Time (Formatted)",
        "Duration (s)",
        "Transcript (Full Verbatim)",
        "QC Valid",
        "QC Errors"
    ]
    writer.writerow(headers)

    for seg in result.segments:
        qc_err_msgs = "; ".join([e.message for e in seg.qc_errors]) if seg.qc_errors else "None"
        writer.writerow([
            result.filename,
            seg.segment_id,
            seg.speaker,
            seg.gender,
            f"{seg.start_time:.3f}",
            f"{seg.end_time:.3f}",
            seg.start_time_str,
            seg.end_time_str,
            f"{seg.duration:.3f}",
            seg.transcript,
            "PASS" if seg.is_valid else "FAIL",
            qc_err_msgs
        ])

    return output.getvalue()


def export_to_tsv(result: TranscriptionResult, delimiter: str = "\t") -> str:
    """Generate TSV string representation."""
    return export_to_csv(result, delimiter=delimiter)


def export_to_txt(result: TranscriptionResult) -> str:
    """Generate clean human-readable TXT transcript."""
    dur = result.audio_info.duration if result.audio_info else (result.segments[-1].end_time if result.segments else 0.0)
    lines = []
    lines.append(f"================================================================================")
    lines.append(f"KARYA TRANSCRIPTION DELIVERABLE: {result.filename}")
    lines.append(f"Language: {result.language} | Script: {result.script}")
    lines.append(f"Duration: {dur}s | Compliance Score: {result.compliance_score}%")
    lines.append(f"================================================================================\n")

    for seg in result.segments:
        lines.append(f"[{seg.start_time_str} --> {seg.end_time_str}] {seg.speaker} ({seg.gender}):")
        lines.append(f"   {seg.transcript}\n")

    return "\n".join(lines)


def export_to_srt(result: TranscriptionResult) -> str:
    """Generate standard SRT subtitle file."""
    def srt_time(seconds: float) -> str:
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds - int(seconds)) * 1000))
        return f"{hrs:02d}:{mins:02d}:{secs:02d},{millis:03d}"

    entries = []
    for i, seg in enumerate(result.segments, 1):
        s_str = srt_time(seg.start_time)
        e_str = srt_time(seg.end_time)
        # SRT standard: subtitle body should be just the transcript text.
        # Speaker label is placed as a clean prefix (compatible with most SRT players).
        speaker_prefix = f"[{seg.speaker}] " if seg.speaker else ""
        entries.append(f"{i}\n{s_str} --> {e_str}\n{speaker_prefix}{seg.transcript}\n")

    return "\n".join(entries)


def export_to_vtt(result: TranscriptionResult) -> str:
    """Generate standard WebVTT subtitle file (.vtt)."""
    def vtt_time(seconds: float) -> str:
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds - int(seconds)) * 1000))
        return f"{hrs:02d}:{mins:02d}:{secs:02d}.{millis:03d}"

    lines = ["WEBVTT", ""]
    for i, seg in enumerate(result.segments, 1):
        s_str = vtt_time(seg.start_time)
        e_str = vtt_time(seg.end_time)
        speaker_prefix = f"[{seg.speaker}] " if seg.speaker else ""
        lines.append(f"{i}")
        lines.append(f"{s_str} --> {e_str}")
        lines.append(f"{speaker_prefix}{seg.transcript}")
        lines.append("")

    return "\n".join(lines)


def export_to_json(result: TranscriptionResult) -> str:
    """Generate clean segment list JSON matching exact Karya user deliverable format."""
    items = []
    for seg in result.segments:
        items.append({
            "start_sec": round(float(seg.start_time), 3),
            "end_sec": round(float(seg.end_time), 3),
            "transcription": seg.transcript or "",
            "speaker": seg.speaker,
            "gender_label": seg.gender
        })
    return json.dumps(items, ensure_ascii=False, indent=2)


def export_to_docx(result: TranscriptionResult, output_path: str) -> str:
    """Generate Word (.docx) audit deliverable."""
    doc = Document()
    
    # Title
    title = doc.add_heading(f"Karya Transcription Deliverable", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Metadata Paragraph
    dur = result.audio_info.duration if result.audio_info else (result.segments[-1].end_time if result.segments else 0.0)
    meta = doc.add_paragraph()
    meta.add_run(f"Audio File: ").bold = True
    meta.add_run(f"{result.filename}\n")
    meta.add_run(f"Target Language & Script: ").bold = True
    meta.add_run(f"{result.language} ({result.script})\n")
    meta.add_run(f"Duration: ").bold = True
    meta.add_run(f"{dur}s | ")
    meta.add_run(f"Compliance Score: ").bold = True
    meta.add_run(f"{result.compliance_score}%\n")
    meta.add_run(f"Total Segments: ").bold = True
    meta.add_run(f"{len(result.segments)} | Total Errors: {result.total_errors}\n")

    # Table
    table = doc.add_table(rows=1, cols=6)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    hdr_cells = table.rows[0].cells
    hdr_titles = ["#", "Start", "End", "Speaker", "Gender", "Verbatim Transcript"]
    for i, t in enumerate(hdr_titles):
        hdr_cells[i].text = t
        for p in hdr_cells[i].paragraphs:
            for run in p.runs:
                run.font.bold = True
                run.font.size = Pt(9.5)

    for seg in result.segments:
        row_cells = table.add_row().cells
        row_cells[0].text = str(seg.segment_id)
        row_cells[1].text = seg.start_time_str
        row_cells[2].text = seg.end_time_str
        row_cells[3].text = seg.speaker
        row_cells[4].text = seg.gender
        row_cells[5].text = seg.transcript
        
        for cell in row_cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(9.0)

    doc.save(output_path)
    return output_path


def export_to_xlsx(result: TranscriptionResult, output_path: str) -> str:
    """Generate formatted Excel (.xlsx) deliverable with summary and data sheets."""
    with pd.ExcelWriter(output_path, engine="openpyxl") as writer:
        # Sheet 1: Data
        rows = []
        for seg in result.segments:
            rows.append({
                "Filename": result.filename,
                "Segment ID": seg.segment_id,
                "Speaker": seg.speaker,
                "Gender": seg.gender,
                "Start Time (s)": seg.start_time,
                "End Time (s)": seg.end_time,
                "Start Time (Formatted)": seg.start_time_str,
                "End Time (Formatted)": seg.end_time_str,
                "Duration (s)": seg.duration,
                "Transcript (Full Verbatim)": seg.transcript,
                "QC Status": "PASS" if seg.is_valid else "FAIL",
                "QC Errors": "; ".join([e.message for e in seg.qc_errors]) if seg.qc_errors else ""
            })
        df_data = pd.DataFrame(rows)
        df_data.to_excel(writer, sheet_name="Transcription", index=False)

        # Sheet 2: Summary
        df_summary = pd.DataFrame([
            {"Metric": "Audio Filename", "Value": result.filename},
            {"Metric": "Language", "Value": result.language},
            {"Metric": "Script", "Value": result.script},
            {"Metric": "Duration (seconds)", "Value": result.audio_info.duration},
            {"Metric": "Sample Rate (Hz)", "Value": result.audio_info.sample_rate},
            {"Metric": "Channels", "Value": result.audio_info.channels},
            {"Metric": "Average RMS (dBFS)", "Value": result.audio_info.rms_db},
            {"Metric": "Estimated SNR (dB)", "Value": result.audio_info.snr_db},
            {"Metric": "Total Segments", "Value": len(result.segments)},
            {"Metric": "Compliance Score (%)", "Value": result.compliance_score},
            {"Metric": "Total QC Errors", "Value": result.total_errors},
            {"Metric": "Total QC Warnings", "Value": result.total_warnings},
            {"Metric": "Is Rejected", "Value": "Yes" if result.is_rejected else "No"},
            {"Metric": "Rejection Reason", "Value": result.rejection_reason or "N/A"},
        ])
        df_summary.to_excel(writer, sheet_name="Audit Summary", index=False)

    return output_path


def export_rejection_csv(rejected_items: List[Dict[str, Any]]) -> str:
    """Generate Rejection Report CSV according to Karya Guideline 3."""
    output = io.StringIO()
    writer = csv.writer(output)
    writer.writerow(["Audio Filename", "Rejection Category", "Detailed Rejection Reason", "Duration (s)", "RMS (dBFS)", "SNR (dB)"])

    for item in rejected_items:
        writer.writerow([
            item.get("filename", ""),
            item.get("rejection_category", ""),
            item.get("rejection_reason", ""),
            item.get("duration", 0.0),
            item.get("rms_db", 0.0),
            item.get("snr_db", 0.0)
        ])

    return output.getvalue()


def export_netflix_srt(events: list) -> str:
    """Generate Netflix-compliant SRT format."""
    def format_time(seconds: float) -> str:
        if seconds is None:
            seconds = 0.0
        seconds = float(seconds)
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds - int(seconds)) * 1000))
        if millis >= 1000:
            millis -= 1000
            secs += 1
        return f"{hrs:02d}:{mins:02d}:{secs:02d},{millis:03d}"

    entries = []
    for i, event in enumerate(events, 1):
        s_val = event.get("start_time", event.get("start", 0.0)) if isinstance(event, dict) else getattr(event, "start_time", getattr(event, "start", 0.0))
        e_val = event.get("end_time", event.get("end", 0.0)) if isinstance(event, dict) else getattr(event, "end_time", getattr(event, "end", 0.0))
        txt_val = event.get("text", "") if isinstance(event, dict) else getattr(event, "text", "")
        
        s_str = format_time(s_val)
        e_str = format_time(e_val)
        text = str(txt_val or "").replace("...", "…")
        entries.append(f"{i}\n{s_str} --> {e_str}\n{text}\n")
    
    return "\n".join(entries)


def export_netflix_vtt(events: list) -> str:
    """Generate Netflix-compliant WebVTT format."""
    def format_time(seconds: float) -> str:
        if seconds is None:
            seconds = 0.0
        seconds = float(seconds)
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds - int(seconds)) * 1000))
        if millis >= 1000:
            millis -= 1000
            secs += 1
        return f"{hrs:02d}:{mins:02d}:{secs:02d}.{millis:03d}"

    lines = ["WEBVTT\n"]
    for i, event in enumerate(events, 1):
        s_val = event.get("start_time", event.get("start", 0.0)) if isinstance(event, dict) else getattr(event, "start_time", getattr(event, "start", 0.0))
        e_val = event.get("end_time", event.get("end", 0.0)) if isinstance(event, dict) else getattr(event, "end_time", getattr(event, "end", 0.0))
        txt_val = event.get("text", "") if isinstance(event, dict) else getattr(event, "text", "")
        
        s_str = format_time(s_val)
        e_str = format_time(e_val)
        text = str(txt_val or "").replace("...", "…")
        lines.append(f"{i}")
        lines.append(f"{s_str} --> {e_str}")
        lines.append(f"{text}")
        lines.append("")
        
    return "\n".join(lines)


def export_netflix_ttml(events: list, language: str = "en") -> str:
    """Generate Netflix TTML/DFXP format."""
    def format_time(seconds: float) -> str:
        if seconds is None:
            seconds = 0.0
        seconds = float(seconds)
        hrs = int(seconds // 3600)
        mins = int((seconds % 3600) // 60)
        secs = int(seconds % 60)
        millis = int(round((seconds - int(seconds)) * 1000))
        if millis >= 1000:
            millis -= 1000
            secs += 1
        return f"{hrs:02d}:{mins:02d}:{secs:02d}.{millis:03d}"

    xml = []
    xml.append('<?xml version="1.0" encoding="utf-8"?>')
    xml.append('<tt xmlns="http://www.w3.org/ns/ttml" xmlns:tts="http://www.w3.org/ns/ttml#styling" xmlns:ttm="http://www.w3.org/ns/ttml#metadata" xmlns:ttp="http://www.w3.org/ns/ttml#parameter" xml:lang="' + language + '">')
    xml.append('  <head>')
    xml.append('    <styling>')
    xml.append('      <style xml:id="default" tts:color="white" tts:fontFamily="sansSerif" tts:fontSize="100%" tts:textAlign="center" tts:origin="10% 10%" tts:extent="80% 80%"/>')
    xml.append('    </styling>')
    xml.append('    <layout>')
    xml.append('      <region xml:id="bottomCenter" style="default" tts:origin="10% 80%" tts:extent="80% 10%"/>')
    xml.append('      <region xml:id="topCenter" style="default" tts:origin="10% 10%" tts:extent="80% 10%"/>')
    xml.append('    </layout>')
    xml.append('  </head>')
    xml.append('  <body>')
    xml.append('    <div region="bottomCenter">')
    
    for event in events:
        s_val = event.get("start_time", event.get("start", 0.0)) if isinstance(event, dict) else getattr(event, "start_time", getattr(event, "start", 0.0))
        e_val = event.get("end_time", event.get("end", 0.0)) if isinstance(event, dict) else getattr(event, "end_time", getattr(event, "end", 0.0))
        txt_val = event.get("text", "") if isinstance(event, dict) else getattr(event, "text", "")
        
        s_str = format_time(s_val)
        e_str = format_time(e_val)
        text = str(txt_val or "").replace("...", "…")
        
        # Basic escaping
        text = text.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")
        
        # Handle italics (which are now escaped)
        text = text.replace("&lt;i&gt;", '<span tts:fontStyle="italic">')
        text = text.replace("&lt;/i&gt;", '</span>')
        text = text.replace("\n", "<br/>")
            
        xml.append(f'      <p begin="{s_str}" end="{e_str}">{text}</p>')
        
    xml.append('    </div>')
    xml.append('  </body>')
    xml.append('</tt>')
    
    return "\n".join(xml)
